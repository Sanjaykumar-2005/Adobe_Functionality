"""
Exercise every tool (except OCR) through the real HTTP endpoints and collect all
outputs into ONE folder, each file named after the functionality that produced it.

Run from the project root:   python verify_all.py
Outputs land in:             ./verification_outputs/
"""
from __future__ import annotations

import io
import json
import os
import platform
import shutil
import subprocess
import sys
import tempfile

import fitz
from PIL import Image
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import app as flask_app
from config import Config
from services import conversion_service

OUT = os.path.join(os.getcwd(), "verification_outputs")
FIX = os.path.join(OUT, "_fixtures")

client = flask_app.app.test_client()
report: list[dict] = []


# --------------------------------------------------------------------------- #
# Fixtures
# --------------------------------------------------------------------------- #
def make_fixtures() -> dict:
    os.makedirs(FIX, exist_ok=True)

    # A 5-page text PDF with a table-ish grid and a heading, so conversions have
    # something real to preserve.
    doc = fitz.open()
    for i in range(5):
        page = doc.new_page()
        page.insert_text((72, 90), f"Section {i + 1}: Quarterly Report",
                         fontname="hebo", fontsize=18)
        page.insert_text((72, 130), "Revenue grew 12% against a flat cost base.",
                         fontname="helv", fontsize=11)
        page.insert_text((72, 150), "Headcount held at 42 across all regions.",
                         fontname="helv", fontsize=11)
        page.draw_rect(fitz.Rect(72, 180, 520, 260), color=(0.2, 0.2, 0.6), width=1)
        page.insert_text((80, 200), "Region    Q1      Q2", fontname="cour", fontsize=10)
        page.insert_text((80, 220), "North     1,204   1,388", fontname="cour", fontsize=10)
        page.insert_text((80, 240), "South       942   1,015", fontname="cour", fontsize=10)
    pdf5 = os.path.join(FIX, "source_5page.pdf")
    doc.save(pdf5)
    doc.close()

    # A second PDF, for merge.
    doc = fitz.open()
    for i in range(2):
        page = doc.new_page()
        page.insert_text((72, 100), f"Appendix page {i + 1}", fontname="helv", fontsize=14)
    pdf2 = os.path.join(FIX, "source_appendix.pdf")
    doc.save(pdf2)
    doc.close()

    # A .docx with a heading, body, and a real table -> Word->PDF via LibreOffice.
    d = Document()
    d.add_heading("Engineering Status", level=1)
    d.add_paragraph("This paragraph should survive the conversion with its styling.")
    t = d.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    data = [["Region", "Q1", "Q2"], ["North", "1204", "1388"], ["South", "942", "1015"]]
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            t.cell(r, c).text = val
    d.add_paragraph("Closing note after the table.")
    docx = os.path.join(FIX, "source_report.docx")
    d.save(docx)

    # Two images -> Image->PDF. DELIBERATELY different pixel sizes AND DPI: when
    # both fixtures were an identical 800x600 the suite could not detect pages of
    # differing size, which is exactly how the "Fit" DPI bug slipped through.
    imgs = []
    for name, color, size, dpi in (
        ("red", (200, 60, 60), (2000, 1500), (300, 300)),   # 300-DPI "scan"
        ("blue", (60, 90, 200), (960, 720), (96, 96)),      # 96-DPI "screenshot"
    ):
        p = os.path.join(FIX, f"source_{name}.png")
        Image.new("RGB", size, color).save(p, dpi=dpi)
        imgs.append(p)

    # A signature image for fill & sign.
    sig = os.path.join(FIX, "source_signature.png")
    Image.new("RGBA", (300, 100), (0, 0, 0, 0)).save(sig)

    return {"pdf5": pdf5, "pdf2": pdf2, "docx": docx, "images": imgs, "sig": sig}


# --------------------------------------------------------------------------- #
# Preflight — the environment checks that decide whether the two LibreOffice
# tools will merely RUN or actually produce the right-looking output.
#
# These are the silent killers on a fresh server: LibreOffice substitutes any
# missing font (layout drifts, no error) and refuses to work from a non-writable
# HOME (emits nothing, no error). Both fail QUIETLY, so check them up front.
# --------------------------------------------------------------------------- #

# Fixtures + real Word docs use Calibri/Cambria. Carlito and Caladea are the
# metric-compatible substitutes; without them LibreOffice picks something of a
# different width and the line breaks move.
NEEDED_FONTS = ("Carlito", "Caladea", "Liberation", "DejaVu")


def _installed_fonts() -> tuple[set[str], str]:
    """Font family names the system can actually see."""
    fc = shutil.which("fc-list")
    if fc:  # Linux / RHEL — the authoritative source for what LibreOffice sees
        try:
            out = subprocess.run([fc, "--format", "%{family}\n"],
                                 capture_output=True, text=True, timeout=30)
            return {line.strip() for line in out.stdout.splitlines() if line.strip()}, "fc-list"
        except (OSError, subprocess.SubprocessError):
            pass
    names = set()
    for d in (r"C:\Windows\Fonts", "/usr/share/fonts", os.path.expanduser("~/.fonts")):
        if os.path.isdir(d):
            for root, _dirs, files in os.walk(d):
                for f in files:
                    if f.lower().endswith((".ttf", ".otf", ".ttc")):
                        names.add(os.path.splitext(f)[0])
    return names, "font directories"


def preflight() -> list[str]:
    """Report environment risks. Returns a list of warnings (empty = all good)."""
    warnings: list[str] = []
    print("=" * 100)
    print("PREFLIGHT")
    print("=" * 100)
    print(f"  platform         : {platform.system()} {platform.release()} "
          f"({platform.machine()})  python {platform.python_version()}")

    # 1. LibreOffice present, and which version (the PDF import filter changes
    #    between versions, so record it alongside the results).
    binary = conversion_service._soffice_binary()
    if binary:
        # On Linux `soffice --version` prints the version; the Windows .exe detaches
        # and prints nothing, so don't present an empty result as if it were real.
        ver = ""
        try:
            out = subprocess.run([binary, "--version"], capture_output=True,
                                 text=True, timeout=60)
            lines = (out.stdout or out.stderr or "").strip().splitlines()
            ver = lines[0].strip() if lines else ""
        except (OSError, subprocess.SubprocessError):
            pass
        print(f"  LibreOffice      : {binary}")
        print(f"                     version: {ver or 'not reported (normal on Windows)'}")
    else:
        print("  LibreOffice      : NOT FOUND")
        warnings.append(
            "LibreOffice is NOT installed. Word->PDF falls back to a TEXT-ONLY "
            "renderer (backgrounds/images/layout lost) and PDF->Word falls back to "
            "the approximate native engine. Install it, or set LIBREOFFICE_BIN.")

    # 2. HOME must be writable, or soffice silently produces nothing.
    home = os.environ.get("HOME") or os.path.expanduser("~")
    writable = False
    try:
        with tempfile.NamedTemporaryFile(dir=home, prefix=".wtest_"):
            writable = True
    except OSError:
        pass
    print(f"  HOME             : {home}  ({'writable' if writable else 'NOT WRITABLE'})")
    if not writable:
        warnings.append(
            f"HOME ({home}) is NOT writable. LibreOffice needs a writable home and "
            "will fail SILENTLY (no output, no error). Set HOME=/tmp for the service "
            "user, as the Dockerfile does.")

    # 3. Fonts — the quiet fidelity killer.
    fonts, how = _installed_fonts()
    missing = [f for f in NEEDED_FONTS
               if not any(f.lower() in name.lower() for name in fonts)]
    found = [f for f in NEEDED_FONTS if f not in missing]
    print(f"  fonts (via {how:<17}): {len(fonts)} families; "
          f"have {found or 'none of the expected'}")
    if missing:
        print(f"                     MISSING: {missing}")
        warnings.append(
            f"Missing font families {missing}. LibreOffice SUBSTITUTES missing fonts "
            "instead of erroring, so conversions will still 'pass' while the layout "
            "silently drifts (wrong line breaks / column widths). On RHEL install: "
            "google-carlito-fonts google-crosextra-caladea-fonts "
            "liberation-fonts dejavu-sans-fonts")

    print()
    return warnings


def upload(path: str):
    with open(path, "rb") as fh:
        return (io.BytesIO(fh.read()), os.path.basename(path))


def run(label: str, url: str, data: dict, expect_files: bool = True) -> dict | None:
    """POST one tool, save its outputs into OUT named after `label`, record result."""
    resp = client.post(url, data=data, content_type="multipart/form-data")
    body = resp.get_json() or {}
    ok = resp.status_code == 200 and body.get("success")

    saved = []
    if ok and expect_files:
        job = body.get("job_id")
        files = body.get("files") or []
        for i, desc in enumerate(files):
            src = os.path.join(Config.OUTPUT_DIR, job, desc["name"])
            if not os.path.exists(src):
                continue
            ext = os.path.splitext(desc["name"])[1]
            suffix = "" if len(files) == 1 else f"_{i + 1}"
            dest = os.path.join(OUT, f"{label}{suffix}{ext}")
            shutil.copy2(src, dest)
            saved.append(os.path.basename(dest))

    report.append({
        "tool": label,
        "http": resp.status_code,
        "ok": bool(ok),
        "engines": body.get("engines"),
        "warning": (body.get("warning") or "")[:60] or None,
        "error": body.get("error"),
        "outputs": saved,
    })
    return body if ok else None


def main() -> int:
    if os.path.isdir(OUT):
        shutil.rmtree(OUT)
    os.makedirs(OUT, exist_ok=True)

    env_warnings = preflight()
    fx = make_fixtures()

    # ---------------- Convert (the two LibreOffice-backed tools) -------------
    run("convert_word-to-pdf", "/api/convert/word-to-pdf",
        {"files": upload(fx["docx"])})

    run("convert_image-to-pdf", "/api/convert/image-to-pdf",
        {"files": [upload(p) for p in fx["images"]], "page_size": "A4", "merge": "true"})

    # "Fit" sizes each page to its image's TRUE physical size (px/dpi*72). The two
    # fixtures are 2000px@300dpi and 960px@96dpi, so the pages SHOULD differ
    # (6.67in vs 10in) — that is Fit working, not the old 1px=1pt bug (which would
    # have made them 27.8in and 13.3in).
    run("convert_image-to-pdf_fit", "/api/convert/image-to-pdf",
        {"files": [upload(p) for p in fx["images"]], "page_size": "Fit", "merge": "true"})

    run("convert_pdf-to-word", "/api/convert/pdf-to-word",
        {"files": upload(fx["pdf5"]), "remove_borders": "false"})

    # ---------------- Organize ----------------------------------------------
    run("organize_merge", "/api/organize/merge",
        {"files": [upload(fx["pdf5"]), upload(fx["pdf2"])]})

    run("organize_split", "/api/organize/split",
        {"file": upload(fx["pdf5"]), "mode": "every_n", "n": "2"})

    run("organize_rotate", "/api/organize/rotate",
        {"file": upload(fx["pdf5"]), "rotation": "90", "pages": "1,3"})

    run("organize_rearrange", "/api/organize/rearrange",
        {"file": upload(fx["pdf5"]), "order": json.dumps([5, 4, 3, 2, 1])})

    run("organize_extract", "/api/organize/extract",
        {"file": upload(fx["pdf5"]), "pages": "2,4"})

    run("organize_delete", "/api/organize/delete",
        {"file": upload(fx["pdf5"]), "pages": "3"})

    run("organize_crop", "/api/organize/crop",
        {"file": upload(fx["pdf5"]),
         "box": json.dumps({"x0": 0.05, "y0": 0.05, "x1": 0.95, "y1": 0.6})})

    run("organize_compress", "/api/organize/compress",
        {"file": upload(fx["pdf5"]), "level": "medium"})

    # ---------------- Edit ---------------------------------------------------
    run("edit_text", "/api/edit/text",
        {"file": upload(fx["pdf5"]),
         "edits": json.dumps([
             {"type": "add_text", "page": 1, "x": 0.1, "y": 0.5,
              "text": "INSERTED BY EDIT TOOL", "font_size": 14, "color": "#cc0000"},
             {"type": "delete_region", "page": 2,
              "x0": 0.1, "y0": 0.1, "x1": 0.9, "y1": 0.2},
         ])})

    run("edit_watermark", "/api/edit/watermark",
        {"file": upload(fx["pdf5"]), "wm_type": "text", "text": "CONFIDENTIAL",
         "opacity": "0.3", "rotation": "45", "font_size": "48",
         "color": "#888888", "position": "center"})

    run("edit_page-numbers", "/api/edit/page-numbers",
        {"file": upload(fx["pdf5"]), "position": "bottom-center", "start": "1",
         "font_size": "12", "color": "#000000", "prefix": "Page ", "suffix": ""})

    run("edit_redact", "/api/edit/redact",
        {"file": upload(fx["pdf5"]),
         "boxes": json.dumps([{"page": 1, "x0": 0.1, "y0": 0.13, "x1": 0.8, "y1": 0.18}])})

    run("edit_fill-sign", "/api/edit/fill-sign",
        {"file": upload(fx["pdf5"]),
         "fields": json.dumps([
             {"type": "text", "page": 1, "x": 0.15, "y": 0.75,
              "text": "Sanjaykumar M", "font_size": 14},
             {"type": "checkbox", "page": 1, "x": 0.15, "y": 0.8, "checked": True},
             {"type": "signature_image", "page": 1,
              "x0": 0.5, "y0": 0.72, "x1": 0.8, "y1": 0.82, "image": "image_0"},
         ]),
         "image_0": upload(fx["sig"])})

    # ---------------- Security ----------------------------------------------
    run("security_protect", "/api/security/protect",
        {"file": upload(fx["pdf5"]), "user_pw": "secret123", "owner_pw": "owner123",
         "print": "true", "modify": "false", "copy": "false", "annotate": "true"})

    # ---------------- Batch --------------------------------------------------
    run("batch_compress", "/api/batch/",
        {"files": [upload(fx["pdf5"]), upload(fx["pdf2"])],
         "operation": "compress", "options": json.dumps({"level": "high"})})

    # ---------------- Report -------------------------------------------------
    print(f"{'TOOL':<26} {'HTTP':<5} {'OK':<4} {'ENGINE':<12} OUTPUT")
    print("-" * 100)
    failures = 0
    for r in report:
        eng = ",".join(r["engines"]) if r["engines"] else "-"
        status = "PASS" if r["ok"] and r["outputs"] else "FAIL"
        if status == "FAIL":
            failures += 1
        out = ", ".join(r["outputs"]) or (r["error"] or "no output")
        print(f"{r['tool']:<26} {r['http']:<5} {status:<4} {eng:<12} {out}")
        if r["warning"]:
            print(f"{'':<26} {'':<5} {'':<4} {'':<12} WARN: {r['warning']}")

    print("-" * 100)
    print(f"{len(report) - failures}/{len(report)} passed. Outputs in: {OUT}")

    # An all-PASS run still means little if the environment is degraded: a missing
    # font never fails a test, it just quietly changes the layout. Say so loudly.
    if env_warnings:
        print()
        print("=" * 100)
        print(f"{len(env_warnings)} ENVIRONMENT WARNING(S) — tools may PASS while the "
              "output is subtly wrong:")
        print("=" * 100)
        for w in env_warnings:
            print(f"  * {w}")

    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
