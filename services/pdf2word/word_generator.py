"""
word_generator.py — the Word Generator stage.

Renders the intermediate document model into a real .docx with python-docx:
sections/page setup, headers/footers, paragraphs with runs + styles + hyperlinks,
headings, lists, TOC rows with dot-leader tab stops, tables, images, and
horizontal rule lines. It contains NO extraction logic — it only draws the model.
"""
from __future__ import annotations

import io
import logging

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from .model import (
    Alignment, Document, Drawing, DrawingKind, Image, ListType, Page,
    Paragraph, Table,
)
from .utils import PT_TO_EMU

log = logging.getLogger("pdf2word.word")

_ALIGN = {
    Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class WordGenerator:
    """Build a .docx from :class:`Page` models, one page at a time."""

    def __init__(self, remove_borders: bool = False) -> None:
        self.doc = DocxDocument()
        self.remove_borders = remove_borders
        self._header_text: str | None = None
        self._footer_text: str | None = None
        self._section_ready = False

    def set_header_footer(self, header: str | None, footer: str | None) -> None:
        self._header_text = header
        self._footer_text = footer

    # ----------------------------------------------------------- page output #
    def add_page(self, page: Page, first: bool) -> None:
        if first:
            self._setup_section(page)
        else:
            self.doc.add_page_break()
        for el in page.elements:
            try:
                if isinstance(el, Paragraph):
                    self._add_paragraph(el, page)
                elif isinstance(el, Table):
                    self._add_table(el)
                elif isinstance(el, Image):
                    self._add_image(el)
                elif isinstance(el, Drawing):
                    self._add_drawing(el)
            except Exception:            # one bad element must not abort the doc
                log.exception("failed to render an element on page %d", page.number)

    def save(self, path: str) -> None:
        self.doc.save(path)
        log.info("DOCX generated: %s", path)

    # --------------------------------------------------------------- section #
    def _setup_section(self, page: Page) -> None:
        s = self.doc.sections[0]
        s.page_width = Emu(int(page.width * PT_TO_EMU))
        s.page_height = Emu(int(page.height * PT_TO_EMU))
        s.orientation = WD_ORIENT.LANDSCAPE if page.width > page.height else WD_ORIENT.PORTRAIT
        s.left_margin = Emu(int(page.margin_left * PT_TO_EMU))
        s.right_margin = Emu(int(page.margin_right * PT_TO_EMU))
        s.top_margin = Emu(int(page.margin_top * PT_TO_EMU))
        s.bottom_margin = Emu(int(page.margin_bottom * PT_TO_EMU))
        if self._header_text:
            s.header.paragraphs[0].text = self._header_text
        if self._footer_text:
            s.footer.paragraphs[0].text = self._footer_text
        self._section_ready = True

    # ------------------------------------------------------------- paragraph #
    def _add_paragraph(self, p: Paragraph, page: Page) -> None:
        style = None
        if p.heading_level:
            style = f"Heading {min(p.heading_level, 4)}"
        elif p.list_type == ListType.BULLET:
            style = "List Bullet"
        elif p.list_type == ListType.NUMBER:
            style = "List Number"

        para = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()

        if p.alignment in _ALIGN:
            para.alignment = _ALIGN[p.alignment]
        pf = para.paragraph_format
        if p.space_before:
            pf.space_before = Pt(min(p.space_before, 240))
        if p.space_after:
            pf.space_after = Pt(min(p.space_after, 240))
        if p.left_indent and not style:
            pf.left_indent = Emu(int(p.left_indent * PT_TO_EMU))

        if p.is_toc:
            self._fill_toc(para, p, page)
            return
        for run in p.runs:
            if not run.text:
                continue
            if run.href:
                self._add_hyperlink(para, run.href, run.text, run.style)
            else:
                self._add_run(para, run.text, run.style)

    def _add_run(self, para, text: str, style) -> None:
        r = para.add_run(text)
        f = r.font
        if style.font:
            f.name = style.font
        if style.size:
            f.size = Pt(style.size)
        f.bold = style.bold
        f.italic = style.italic
        f.underline = style.underline
        try:
            f.color.rgb = RGBColor(*style.color)
        except Exception:
            pass

    def _add_hyperlink(self, para, url: str, text: str, style) -> None:
        r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rPr.append(col)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        if style and style.size:
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(style.size * 2))); rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        run.append(t)
        link.append(run)
        para._p.append(link)

    def _fill_toc(self, para, p: Paragraph, page: Page) -> None:
        """Render a TOC row: title + right-aligned page number with a dot leader."""
        content_w = page.width - page.margin_left - page.margin_right
        pos = Emu(int(max(72.0, content_w) * PT_TO_EMU))
        para.paragraph_format.tab_stops.add_tab_stop(
            pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        title_style = p.runs[0].style if p.runs else None
        self._add_run(para, p.text, title_style) if p.runs else para.add_run(p.text)
        para.add_run("\t")
        para.add_run(p.toc_page or "")

    # ----------------------------------------------------------------- table #
    def _add_table(self, t: Table) -> None:
        rows = t.rows
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=ncols)
        if t.has_borders and not self.remove_borders:
            try:
                table.style = "Table Grid"
            except Exception:
                pass
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                table.rows[ri].cells[ci].text = row[ci] if ci < len(row) else ""
        if self.remove_borders or not t.has_borders:
            self._strip_table_borders(table)

    @staticmethod
    def _strip_table_borders(table) -> None:
        tblPr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); borders.append(e)
        tblPr.append(borders)

    # ----------------------------------------------------------------- image #
    def _add_image(self, img: Image) -> None:
        if not img.data:
            return
        width = Emu(int(max(1.0, img.width) * PT_TO_EMU))
        self.doc.add_picture(io.BytesIO(img.data), width=width)

    # -------------------------------------------------------------- drawings #
    def _add_drawing(self, d: Drawing) -> None:
        if d.kind not in (DrawingKind.HLINE,):
            return  # vertical lines / rects have no clean flowing equivalent
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        pPr = para._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(max(4, int((d.width or 1) * 4))))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "%02X%02X%02X" % (d.color or (0, 0, 0)))
        pbdr.append(bottom)
        pPr.append(pbdr)
