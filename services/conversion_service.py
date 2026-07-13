"""
Format-conversion service: Word <-> PDF and Image -> PDF.

Pure Python (no Flask). Functions take absolute input paths + plain params,
write outputs under the job OUTPUT dir, and return lists of absolute paths.

External engines are imported defensively. None of them are required to import
this module; a missing engine surfaces as a clear :class:`RuntimeError` only when
the relevant function is actually called and no fallback succeeds.
"""
from __future__ import annotations

import io
import os
import pathlib
import shutil
import subprocess
import tempfile
import threading

import fitz

from config import Config
from utils.file_utils import output_path, stem, with_suffix

# ---- Optional engines (guarded) ------------------------------------------- #
try:
    from PIL import Image  # type: ignore
    _PIL_OK = True
except Exception:  # pragma: no cover
    Image = None  # type: ignore
    _PIL_OK = False

try:
    from pdf2docx import Converter  # type: ignore
    _PDF2DOCX_OK = True
except Exception:  # pragma: no cover
    Converter = None  # type: ignore
    _PDF2DOCX_OK = False


# --------------------------------------------------------------------------- #
# LibreOffice — the primary conversion engine
# --------------------------------------------------------------------------- #
# Every LibreOffice conversion runs headless through the `soffice` binary. It is
# the preferred engine for BOTH directions (Word->PDF and PDF->Word) because it
# uses real office rendering/import filters, needs no Windows and no MS Word, and
# works the same on a Linux server as on a dev laptop (see the Dockerfile).
_SOFFICE_TIMEOUT = int(os.environ.get("SOFFICE_TIMEOUT", "300"))

# A headless soffice costs a few hundred MB of RSS, and nothing else bounds how
# many run at once: gunicorn serves workers*threads requests concurrently and each
# conversion forks its own process. Left unbounded that is the app's largest
# memory spike and the usual reason the container gets OOM-killed. Queue them.
_SOFFICE_MAX_CONCURRENCY = max(1, int(os.environ.get("SOFFICE_MAX_CONCURRENCY", "1")))
_SOFFICE_SLOTS = threading.Semaphore(_SOFFICE_MAX_CONCURRENCY)

_SOFFICE_WIN_PATHS = (
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
)


def _soffice_binary() -> str | None:
    """Locate the LibreOffice headless binary, if installed.

    Honours ``LIBREOFFICE_BIN`` (or ``SOFFICE_BIN``) so a non-standard install —
    or the one baked into the Docker image — can be pointed at explicitly.
    """
    for var in ("LIBREOFFICE_BIN", "SOFFICE_BIN"):
        explicit = os.environ.get(var)
        if explicit and os.path.exists(explicit):
            return explicit
    for name in ("soffice", "libreoffice", "soffice.exe"):
        found = shutil.which(name)
        if found:
            return found
    for path in _SOFFICE_WIN_PATHS:
        if os.path.exists(path):
            return path
    return None


def soffice_available() -> bool:
    """True when LibreOffice can be used (callers pick their fallback on this)."""
    return _soffice_binary() is not None


def _run_soffice(args: list[str], out_dir: str) -> bool:
    """Run one headless LibreOffice conversion. True when the process succeeded.

    Each call gets a PRIVATE user profile (``-env:UserInstallation``). Without it,
    two conversions running at once share ``~/.config/libreoffice`` and the second
    one either blocks on the profile lock or exits doing nothing — which on a
    multi-user server shows up as random empty outputs. The profile is a temp dir,
    removed afterwards.
    """
    binary = _soffice_binary()
    if not binary:
        return False

    # Bounded wait: a caller that cannot get a slot in time gives up rather than
    # piling up behind the queue until gunicorn's own timeout kills the worker.
    if not _SOFFICE_SLOTS.acquire(timeout=_SOFFICE_TIMEOUT):
        return False

    profile = tempfile.mkdtemp(prefix="lo_profile_")
    cmd = [
        binary,
        f"-env:UserInstallation={pathlib.Path(profile).as_uri()}",
        "--headless", "--norestore", "--invisible", "--nolockcheck", "--nodefault",
        *args,
        "--outdir", out_dir,
    ]
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=_SOFFICE_TIMEOUT)
        return True
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
        return False
    finally:
        _SOFFICE_SLOTS.release()
        shutil.rmtree(profile, ignore_errors=True)


def _soffice_convert(src: str, out_dir: str, convert_to: str, ext: str,
                     infilter: str | None = None) -> str | None:
    """Convert *src* with LibreOffice; return the produced path (or ``None``).

    LibreOffice always names its output ``<stem-of-src>.<ext>`` inside *out_dir*.
    """
    args = []
    if infilter:
        args.append(f"--infilter={infilter}")
    args += ["--convert-to", convert_to, src]
    if not _run_soffice(args, out_dir):
        return None
    produced = os.path.join(out_dir, stem(src) + ext)
    return produced if os.path.exists(produced) else None


# --------------------------------------------------------------------------- #
# Word -> PDF
# --------------------------------------------------------------------------- #
def _word_to_pdf_libreoffice(src: str, out_dir: str) -> str | None:
    """Convert .docx/.doc -> PDF with LibreOffice. Returns the PDF path or ``None``."""
    return _soffice_convert(src, out_dir, "pdf", ".pdf")


def _word_to_pdf_word_com(src: str, dest: str) -> str | None:
    """Convert via Microsoft Word using COM automation (Windows + Word only).

    Driven through PowerShell, so NO extra Python package (docx2pdf / pywin32) is
    required — only Windows with Word installed. This uses Word's own rendering
    engine, so the PDF is a FULL-FIDELITY copy of the .docx (backgrounds, images,
    fonts, layout). Returns the produced PDF path, or ``None`` if unavailable.
    """
    if os.name != "nt":
        return None
    # Escape single quotes for the single-quoted PowerShell string literals.
    s = src.replace("'", "''")
    d = dest.replace("'", "''")
    ps = (
        "$ErrorActionPreference='Stop';"
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        "$doc=$null;"
        "try{"
        f"$doc=$word.Documents.Open('{s}',$false,$true);"   # (path, ConfirmConversions=false, ReadOnly=true)
        f"$doc.ExportAsFixedFormat('{d}',17);"               # 17 = wdExportFormatPDF
        "}finally{"
        "if($doc){$doc.Close($false)};"
        "$word.Quit()"
        "}"
    )
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", ps],
            check=True, capture_output=True, timeout=180,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
        return None
    return dest if os.path.exists(dest) else None


def _word_to_pdf_docx2pdf(src: str, dest: str) -> str | None:
    """Convert via the docx2pdf package (needs MS Word, Windows/macOS)."""
    try:
        from docx2pdf import convert as _convert  # type: ignore
    except Exception:
        return None
    try:
        _convert(src, dest)
    except Exception:
        return None
    return dest if os.path.exists(dest) else None


def _word_to_pdf_reportlab(src: str, dest: str) -> str:
    """Pure-Python best-effort fallback: python-docx -> reportlab.

    Renders paragraphs (with heading sizing + basic bold/italic) and simple
    tables. Always produces a PDF. Raises :class:`RuntimeError` only if the
    pure-python toolchain itself is missing.
    """
    try:
        from docx import Document  # python-docx
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "Word->PDF fallback needs python-docx. Install LibreOffice or "
            "python-docx + reportlab."
        ) from exc
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
        )
        from reportlab.lib import colors
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "Word->PDF fallback needs reportlab. Install LibreOffice or reportlab."
        ) from exc

    docx = Document(src)
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    flow = []

    def _inline(par) -> str:
        """Reassemble a paragraph's runs with crude bold/italic markup."""
        parts = []
        for run in par.runs:
            txt = (run.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if not txt:
                continue
            if run.bold:
                txt = f"<b>{txt}</b>"
            if run.italic:
                txt = f"<i>{txt}</i>"
            parts.append(txt)
        return "".join(parts) or (par.text or "")

    for par in docx.paragraphs:
        text = _inline(par)
        if not text.strip():
            flow.append(Spacer(1, 6))
            continue
        name = (par.style.name or "").lower() if par.style else ""
        if name.startswith("heading") or name == "title":
            size = 18 if "1" in name or name == "title" else 14
            style = ParagraphStyle("h", parent=body, fontSize=size,
                                   leading=size + 4, spaceAfter=8, spaceBefore=8)
            flow.append(Paragraph(text, style))
        else:
            flow.append(Paragraph(text, body))

    for table in docx.tables:
        data = [[(cell.text or "") for cell in row.cells] for row in table.rows]
        if not data:
            continue
        tbl = Table(data)
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        flow.append(Spacer(1, 6))
        flow.append(tbl)

    if not flow:
        flow.append(Paragraph("(empty document)", body))

    SimpleDocTemplate(dest, pagesize=A4,
                      leftMargin=inch, rightMargin=inch,
                      topMargin=inch, bottomMargin=inch).build(flow)
    return dest


def word_to_pdf(paths: list[str], job_id: str,
                engines_out: list[str] | None = None) -> list[str]:
    """Convert each .docx/.doc to PDF, trying LibreOffice -> docx2pdf -> reportlab.

    Returns one PDF per input (in order). The pure-python fallback guarantees a
    result even with no office suite installed — but it is TEXT-ONLY (drops
    backgrounds, images, colors, complex layout).

    If ``engines_out`` is provided, the engine used for each file is appended to
    it ("libreoffice" / "word" / "reportlab") so callers can warn on fallback.
    """
    if not paths:
        raise ValueError("No input documents provided.")
    out_dir = os.path.dirname(output_path(job_id, "x.pdf"))
    results: list[str] = []
    for src in paths:
        dest = output_path(job_id, with_suffix(src, "", ".pdf"))
        engine = "libreoffice"
        produced = _word_to_pdf_libreoffice(src, out_dir)
        if produced and os.path.abspath(produced) != os.path.abspath(dest):
            shutil.move(produced, dest)
            produced = dest
        if not produced:
            engine = "word"
            produced = _word_to_pdf_word_com(src, dest)
        if not produced:
            engine = "word"
            produced = _word_to_pdf_docx2pdf(src, dest)
        if not produced:
            engine = "reportlab"
            produced = _word_to_pdf_reportlab(src, dest)
        results.append(produced)
        if engines_out is not None:
            engines_out.append(engine)
    return results


# --------------------------------------------------------------------------- #
# Image -> PDF
# --------------------------------------------------------------------------- #
def _prep_image(path: str):
    """Open an image and flatten transparency/palette onto white RGB.

    Carries the source DPI across the flatten: ``Image.new`` starts with an empty
    ``info``, so without this a transparent PNG would lose its resolution and the
    "Fit" page size below would fall back to the default DPI.
    """
    im = Image.open(path)
    if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
        rgba = im.convert("RGBA")
        bg = Image.new("RGB", rgba.size, (255, 255, 255))
        bg.paste(rgba, mask=rgba.split()[-1])
        if "dpi" in im.info:
            bg.info["dpi"] = im.info["dpi"]
        return bg
    return im.convert("RGB")


# Images with no DPI metadata (screenshots, most web images) are assumed to be
# screen resolution. 96 is the near-universal convention.
_DEFAULT_DPI = 96.0


def _image_dpi(im) -> float:
    """The image's horizontal DPI, or a sane default.

    Metadata is routinely absent or junk (0, or absurd values from bad scanners),
    so anything implausible is rejected rather than trusted into a broken page size.
    """
    dpi = im.info.get("dpi")
    if isinstance(dpi, (tuple, list)) and dpi:
        dpi = dpi[0]
    try:
        value = float(dpi)
    except (TypeError, ValueError):
        return _DEFAULT_DPI
    return value if 1.0 <= value <= 2400.0 else _DEFAULT_DPI


def _fit_page_size(im) -> tuple[float, float]:
    """Page size in POINTS for a "Fit" page: the image's true physical size.

    The old code treated 1 pixel as 1 point, which ignores DPI entirely — a
    2000px-wide 300-DPI scan became a 27.8in page instead of 6.67in. Points are
    1/72in, so the honest conversion is ``px / dpi * 72``.
    """
    dpi = _image_dpi(im)
    return im.width * 72.0 / dpi, im.height * 72.0 / dpi


def _place_on_page(im, page_size):
    """Return an RGB page-sized canvas with *im* centered, aspect preserved.

    ``page_size`` is a ``(w, h)`` point tuple. "Fit" pages do not go through here —
    they are written directly at their true physical size (see ``_save_fit_pdf``).
    """
    pw, ph = page_size
    canvas = Image.new("RGB", (int(pw), int(ph)), (255, 255, 255))
    scale = min(pw / im.width, ph / im.height)
    new_w, new_h = max(1, int(im.width * scale)), max(1, int(im.height * scale))
    resized = im.resize((new_w, new_h), Image.LANCZOS)
    canvas.paste(resized, ((int(pw) - new_w) // 2, (int(ph) - new_h) // 2))
    return canvas


def _save_fit_pdf(images: list, dest: str) -> None:
    """Write "Fit" pages, each at its image's true physical size.

    Pillow cannot do this: its PDF writer applies ONE global ``resolution`` to every
    page and ignores per-image ``info["dpi"]`` (verified on Pillow 12), so a
    multi-image save would force every page to the same scale. PyMuPDF lets each page
    carry its own size, and the image is embedded at FULL resolution — no resampling.
    """
    doc = fitz.open()
    try:
        for im in images:
            pw, ph = _fit_page_size(im)
            page = doc.new_page(width=pw, height=ph)
            buf = io.BytesIO()
            im.save(buf, "PNG")
            page.insert_image(fitz.Rect(0, 0, pw, ph), stream=buf.getvalue())
        doc.save(dest)
    finally:
        doc.close()


def image_to_pdf(paths: list[str], job_id: str, *, page_size: str = "A4",
                 merge: bool = True, out_name: str = "images.pdf") -> list[str]:
    """Convert images to PDF using Pillow.

    *page_size* is a key of ``Config.PAGE_SIZES`` ("Fit" makes each page match its
    image). Images keep aspect ratio, centered on white. ``merge=True`` yields one
    multi-page PDF; otherwise one PDF per image.
    """
    if not _PIL_OK:
        raise RuntimeError("Image->PDF requires Pillow, which is not installed.")
    if not paths:
        raise ValueError("No input images provided.")
    if page_size not in Config.PAGE_SIZES:
        raise ValueError(
            f"Unknown page_size '{page_size}'. "
            f"Choose from: {', '.join(Config.PAGE_SIZES)}."
        )
    size = Config.PAGE_SIZES[page_size]
    images = [_prep_image(p) for p in paths]

    # "Fit" (size is None): every page takes its own image's physical size, so the
    # pages legitimately differ. It cannot go through Pillow's writer, which forces
    # a single resolution on all pages.
    if size is None:
        if merge:
            dest = output_path(job_id, out_name)
            _save_fit_pdf(images, dest)
            return [dest]
        outputs: list[str] = []
        for src, im in zip(paths, images):
            dest = output_path(job_id, with_suffix(src, "", ".pdf"))
            _save_fit_pdf([im], dest)
            outputs.append(dest)
        return outputs

    pages = [_place_on_page(im, size) for im in images]

    if merge:
        dest = output_path(job_id, out_name)
        first, rest = pages[0], pages[1:]
        first.save(dest, "PDF", resolution=72.0, save_all=True, append_images=rest)
        return [dest]

    outputs = []
    for src, page in zip(paths, pages):
        dest = output_path(job_id, with_suffix(src, "", ".pdf"))
        page.save(dest, "PDF", resolution=72.0)
        outputs.append(dest)
    return outputs


# --------------------------------------------------------------------------- #
# PDF -> Word
# --------------------------------------------------------------------------- #
def _pdf_to_word_libreoffice(src: str, dest: str) -> str | None:
    """Convert a PDF to .docx with LibreOffice's PDF import filter.

    ``--infilter=writer_pdf_import`` forces the PDF into **Writer** (the default
    would open it in Draw, which cannot save .docx at all). LibreOffice then lays
    the page out faithfully — backgrounds, borders, shading, images and colours all
    survive, which is what the native rule-based engine cannot do.

    TRADE-OFF, and it is a real one: Writer's PDF import places text in absolutely
    positioned FRAMES, not flowing paragraphs. The document *looks* like the
    original but is awkward to edit or reflow in Word. Callers who need genuinely
    editable text should prefer ``services.pdf2word`` instead.

    Returns the produced .docx path, or ``None`` if LibreOffice is unavailable or
    the conversion failed (caller then falls back).
    """
    out_dir = os.path.dirname(dest)
    produced = _soffice_convert(
        src, out_dir, "docx:MS Word 2007 XML", ".docx", infilter="writer_pdf_import",
    )
    if not produced:
        return None
    if os.path.abspath(produced) != os.path.abspath(dest):
        shutil.move(produced, dest)
    return dest


def word_is_running() -> bool:
    """True if Microsoft Word (WINWORD.EXE) is currently running (Windows only).

    The faithful PDF->Word path (``_pdf_to_word_word_com``) dismisses Word's "convert
    PDF" dialog by activating the Word window and sending Enter. If the user already
    has Word open, those keystrokes hit the wrong window, so the faithful conversion
    fails and silently falls back to the lower-fidelity pdf2docx path. Callers use this
    to tell the user to close Word first. Always ``False`` off Windows (no Word path).
    """
    if os.name != "nt":
        return False
    try:
        out = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE", "/NH"],
            capture_output=True, text=True, timeout=10,
        )
        return "WINWORD.EXE" in (out.stdout or "")
    except Exception:
        return False


def _pdf_to_word_word_com(src: str, dest: str) -> str | None:
    """Convert a PDF to .docx using Microsoft Word's own PDF import (Windows + Word).

    Word reconstructs the document faithfully (all columns/rows, page framing, no
    right-shift) — far better than pdf2docx. The obstacle on this machine: opening a
    PDF in Word raises a modal "Word will now convert your PDF" dialog that freezes
    headless automation. We work around it by opening the PDF in a VISIBLE Word (in a
    background job) while a foreground loop repeatedly activates the Word window and
    presses Enter to dismiss that dialog. NO registry/security change is needed.

    Trade-off: Word appears briefly and grabs keyboard focus for a few seconds while
    converting — so this is best on a local/desktop run, not a shared server. Returns
    the produced .docx path, or ``None`` if Word is unavailable / it times out (caller
    then falls back to pdf2docx).
    """
    if os.name != "nt":
        return None
    s = src.replace("'", "''")
    d = dest.replace("'", "''")
    script = _WORD_PDF_PS1.replace("__SRC__", s).replace("__DEST__", d)
    ps1 = dest + ".convert.ps1"
    try:
        with open(ps1, "w", encoding="utf-8") as fh:
            fh.write(script)
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", ps1],
            check=False, capture_output=True, timeout=240,
        )
    except (subprocess.TimeoutExpired, OSError):
        return None
    finally:
        try:
            os.remove(ps1)
        except OSError:
            pass
    return dest if os.path.exists(dest) else None


# PowerShell used by _pdf_to_word_word_com. Opens the PDF in a visible Word (in a
# background job) and, from the foreground, keeps pressing Enter on Word's modal
# "convert PDF" dialog until the .docx is written. Cleans up any invisible orphan
# Word instances at the end.
_WORD_PDF_PS1 = r"""
$src  = '__SRC__'
$dest = '__DEST__'
if (Test-Path $dest) { Remove-Item $dest -Force -ErrorAction SilentlyContinue }
$job = Start-Job -ScriptBlock {
    param($src, $dest)
    try {
        $w = New-Object -ComObject Word.Application
        $w.Visible = $true
        $w.DisplayAlerts = 0
        $d = $w.Documents.Open($src, $false, $true)
        $d.SaveAs2($dest, 16)
        $d.Close($false)
        $w.Quit()
    } catch {}
} -ArgumentList $src, $dest
Add-Type -AssemblyName System.Windows.Forms
$sh = New-Object -ComObject WScript.Shell
for ($i = 0; $i -lt 60; $i++) {
    Start-Sleep -Milliseconds 700
    try {
        if ($sh.AppActivate('Microsoft Word')) {
            Start-Sleep -Milliseconds 150
            [System.Windows.Forms.SendKeys]::SendWait('~')
        }
    } catch {}
    if (Test-Path $dest) { break }
}
Wait-Job $job -Timeout 5 | Out-Null
Stop-Job $job -ErrorAction SilentlyContinue
Remove-Job $job -Force -ErrorAction SilentlyContinue
Get-Process WINWORD -ErrorAction SilentlyContinue |
    Where-Object { [string]::IsNullOrWhiteSpace($_.MainWindowTitle) } |
    Stop-Process -Force -ErrorAction SilentlyContinue
"""


def _repair_docx_table_layout(path: str) -> None:
    """Fix pdf2docx tables whose column grid doesn't match their cell widths.

    pdf2docx frequently emits a ``<w:tblGrid>`` of equal columns (e.g. three 2.76in
    thirds) while the cells declare very different widths (e.g. 0.34in / 7.59in /
    0.34in) and leaves ``tblW="auto"``. Word then lays the table out on the *grid*,
    so a wide content cell starts on the wrong column — the whole block shifts right
    and overflows the page (the classic "converted document is pushed to the right").
    Because pdf2docx wraps each page's content in such layout tables (often nested),
    this hits richly-formatted templates hard.

    For every table (including nested ones) whose grid disagrees with its first row's
    real cell widths, we: rewrite the grid to match the cells, set an explicit
    ``tblW`` = sum of cell widths, force ``tblLayout=fixed`` (so Word stops
    auto-fitting onto the bad grid), and zero any stray table indent. Tables whose
    grid already matches (normal Word data tables) are left untouched. Best-effort:
    never raises.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def tw(v):
        try:
            return int(float(v))
        except (TypeError, ValueError):
            return None

    try:
        doc = Document(path)
        body = doc.element.body
        changed = False
        for tbl in body.iter(qn("w:tbl")):
            tblPr = tbl.find(qn("w:tblPr"))
            grid = tbl.find(qn("w:tblGrid"))
            tr = tbl.find(qn("w:tr"))
            if tblPr is None or grid is None or tr is None:
                continue
            # Intended column widths from the first row's cells (twips), honoring gridSpan.
            cell_w, ok = [], True
            for tc in tr.findall(qn("w:tc")):
                tcPr = tc.find(qn("w:tcPr"))
                span, w = 1, None
                if tcPr is not None:
                    gs = tcPr.find(qn("w:gridSpan"))
                    if gs is not None:
                        try:
                            span = max(1, int(gs.get(qn("w:val"))))
                        except (TypeError, ValueError):
                            span = 1
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None and tcW.get(qn("w:type")) in (None, "dxa"):
                        w = tw(tcW.get(qn("w:w")))
                if w is None:
                    ok = False
                    break
                per = int(round(w / span))
                cell_w.extend([per] * span)
            cols = grid.findall(qn("w:gridCol"))
            if not ok or not cell_w or len(cols) != len(cell_w):
                continue
            cur = [tw(c.get(qn("w:w"))) for c in cols]
            if cur == cell_w:
                continue  # grid already consistent — leave normal tables alone
            for c, nw in zip(cols, cell_w):
                c.set(qn("w:w"), str(nw))
            total = sum(cell_w)
            tblW = tblPr.find(qn("w:tblW"))
            if tblW is None:
                tblW = OxmlElement("w:tblW")
                tblPr.append(tblW)
            tblW.set(qn("w:w"), str(total))
            tblW.set(qn("w:type"), "dxa")
            lay = tblPr.find(qn("w:tblLayout"))
            if lay is None:
                lay = OxmlElement("w:tblLayout")
                tblPr.append(lay)
            lay.set(qn("w:type"), "fixed")
            ind = tblPr.find(qn("w:tblInd"))
            if ind is not None:
                ind.set(qn("w:w"), "0")
                ind.set(qn("w:type"), "dxa")
            changed = True
        if changed:
            doc.save(path)
    except Exception:
        return


def _normalize_docx_left_shift(path: str) -> None:
    """Undo the uniform right-shift pdf2docx adds to a converted .docx.

    pdf2docx preserves each block's absolute x-position by adding a left indent to
    every paragraph/table, which makes the whole page look shifted to the right.

    We estimate the uniform shift as the MOST COMMON positive left-indent across
    all text (body paragraphs *and* paragraphs inside tables — pdf2docx often wraps
    content in invisible layout tables), then subtract it from every paragraph and
    table. Content clamped below zero snaps back to the page margin, so lines that
    were already at the margin stay put while intentional *relative* indentation
    (bullets, quotes, deeper nesting) is preserved. No-op unless a real chunk of the
    text shares the same positive indent, so clean docs are untouched. Best-effort:
    never raises (a failure just leaves the doc as-is).
    """
    try:
        from collections import Counter
        from docx import Document
        from docx.shared import Emu
        from docx.oxml.ns import qn
    except Exception:
        return

    def emu(v) -> int:
        return int(v) if v is not None else 0

    def all_paragraphs(doc):
        # Body paragraphs + every table-cell paragraph (covers layout tables).
        yield from doc.paragraphs
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    try:
        doc = Document(path)
        indents = [emu(p.paragraph_format.left_indent)
                   for p in all_paragraphs(doc) if (p.text or "").strip()]
        pos = [i for i in indents if i > 45720]  # ignore < ~0.05in
        if not pos:
            return
        # Bucket to ~0.02in and take the most common indent as the uniform shift.
        bucket = 18288  # 0.02in in EMU
        counts = Counter(round(i / bucket) * bucket for i in pos)
        delta, freq = counts.most_common(1)[0]
        # Only correct when the shift affects a real share of the text.
        if delta <= 45720 or freq < max(2, len(pos) * 0.3):
            return
        for p in all_paragraphs(doc):
            new = emu(p.paragraph_format.left_indent) - delta
            p.paragraph_format.left_indent = Emu(new) if new > 0 else None
        # Shift tables by the same amount (w:tblInd is in twips; 1in = 1440 twips).
        delta_tw = int(round(delta / 635))
        for tbl in doc.tables:
            tblPr = tbl._tbl.tblPr
            ind = tblPr.find(qn("w:tblInd")) if tblPr is not None else None
            if ind is None:
                continue
            try:
                w = int(ind.get(qn("w:w")))
            except (TypeError, ValueError):
                w = 0
            ind.set(qn("w:w"), str(max(0, w - delta_tw)))
            ind.set(qn("w:type"), "dxa")
        doc.save(path)
    except Exception:
        return


def _strip_docx_table_borders(path: str) -> None:
    """Make every table in a .docx borderless (remove all table + cell borders).

    Why: a PDF has no "borderless table" concept — it stores only the visual result.
    When a PDF is imported back to Word, the importer reconstructs each detected table
    and stamps a border on every cell, so tables that were borderless in the original
    "come up" with full gridlines. This removes those borders again for users who want
    the clean, borderless look back.

    Caveat: this is all-or-nothing — it also removes borders from tables that were
    *meant* to have them (the PDF can't tell the two apart), so it's opt-in. Walks all
    tables incl. nested ones. Best-effort: never raises.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    edges = ("top", "left", "bottom", "right", "insideH", "insideV")

    def _none_borders(tag: str):
        el = OxmlElement(tag)
        for edge in edges:
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            e.set(qn("w:space"), "0")
            el.append(e)
        return el

    try:
        doc = Document(path)
        body = doc.element.body
        changed = False
        for tbl in body.iter(qn("w:tbl")):
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)
            old = tblPr.find(qn("w:tblBorders"))
            if old is not None:
                tblPr.remove(old)
            tblPr.append(_none_borders("w:tblBorders"))
            # Clear any per-cell borders so they don't override the table setting.
            for tc in tbl.iter(qn("w:tc")):
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    continue
                cb = tcPr.find(qn("w:tcBorders"))
                if cb is not None:
                    tcPr.remove(cb)
            changed = True
        if changed:
            doc.save(path)
    except Exception:
        return


def pdf_to_word(paths: list[str], job_id: str,
                engines_out: list[str] | None = None,
                remove_borders: bool = False) -> list[str]:
    """Convert each PDF to .docx (one per input).

    Fidelity strategy (best first):

    1. **LibreOffice** (``_pdf_to_word_libreoffice``) — the primary engine. Runs
       anywhere (Linux server, Docker, Windows), needs no MS Word, and preserves
       backgrounds / borders / shading / images. Its text lands in positioned
       frames, so the result is faithful to look at but awkward to re-edit.
    2. **Microsoft Word's own PDF import** (``_pdf_to_word_word_com``) — Windows +
       Word only, and only reached when LibreOffice is not installed. Briefly opens
       Word and steals keyboard focus, so it is a desktop-only path.
    3. **pdf2docx fallback** — pure-Python, no office suite needed. A best-effort
       *guess* at the layout (a PDF has no logical structure), so on richly
       formatted templates it can shift content and drop hard-to-parse cells. We
       post-process it with ``_repair_docx_table_layout`` (fixes the right-shift /
       overflow from pdf2docx's broken table grids) and ``_normalize_docx_left_shift``
       (residual paragraph indent).

    If ``engines_out`` is provided, the engine used for each file is appended to it
    ("libreoffice" / "word" / "pdf2docx") so callers can warn the user when a file
    did NOT get a faithful conversion.

    ``remove_borders=True`` strips all table/cell borders from each output (see
    ``_strip_docx_table_borders``) — for users who want borderless tables back after
    the PDF importer stamped gridlines on every reconstructed cell.

    Raises :class:`RuntimeError` only if EVERY path is unavailable, or the last-resort
    conversion itself fails.
    """
    if not paths:
        raise ValueError("No input PDFs provided.")
    outputs: list[str] = []
    for src in paths:
        dest = output_path(job_id, with_suffix(src, "", ".docx"))

        # 1) Preferred: LibreOffice (works on the server; no Word, no focus stealing).
        produced = _pdf_to_word_libreoffice(src, dest)
        engine = "libreoffice"

        # 2) Only when LibreOffice is absent: Word's own importer (Windows desktop).
        if not produced:
            produced = _pdf_to_word_word_com(src, dest)
            engine = "word"

        if produced:
            if remove_borders:
                _strip_docx_table_borders(produced)
            outputs.append(produced)
            if engines_out is not None:
                engines_out.append(engine)
            continue

        # 3) Last resort: pdf2docx + layout repairs.
        if not _PDF2DOCX_OK:
            raise RuntimeError(
                "PDF->Word needs LibreOffice (install it, or set LIBREOFFICE_BIN), "
                "Microsoft Word (Windows), or the 'pdf2docx' package — none of which "
                "are available."
            )
        try:
            conv = Converter(src)
            try:
                conv.convert(dest)
            finally:
                conv.close()
        except Exception as exc:
            raise RuntimeError(
                f"Failed to convert '{os.path.basename(src)}' to Word: {exc}"
            ) from exc
        _repair_docx_table_layout(dest)   # fix pdf2docx's broken table grids (right-shift/overflow)
        _normalize_docx_left_shift(dest)  # undo any residual uniform paragraph indent shift
        if remove_borders:
            _strip_docx_table_borders(dest)
        outputs.append(dest)
        if engines_out is not None:
            engines_out.append("pdf2docx")
    return outputs
